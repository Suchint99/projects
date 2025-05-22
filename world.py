from docx import Document
from docx.shared import Pt

# Placeholder for AI content generation (use an AI API or model here)
def generate_summary(topic):
    # Imagine this function calls an AI API to generate content based on the topic
    return f"This is an AI-generated summary on the topic of {topic}."

# Function to create a Word document with AI-generated content
def create_word_document(topics):
    # Create a new Document
    doc = Document()
    doc.add_heading("AI-Generated Document", 0)

    # Loop through each topic, generate AI content, and add it to the document
    for topic in topics:
        # Generate content using AI
        content = generate_summary(topic)

        # Add topic as a heading
        doc.add_heading(topic, level=1)

        # Add content as a paragraph
        paragraph = doc.add_paragraph(content)
        paragraph.style.font.size = Pt(12)  # Set font size if needed

    # Save the document
    doc.save("AI_Generated_World_File.docx")
    print("Word document created successfully.")

# Example usage
topics = ["Climate Change", "Artificial Intelligence", "Space Exploration"]
create_word_document(topics)

